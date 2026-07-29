import os
import json
import logging
import re
import time
import threading
from collections import deque
import requests
import pdfplumber
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
import docx
import config
from experience_calculator import safe_process_record, dedupe_publications, renumber_publications

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Per-model RPM limits (requests per minute). Keys accept both provider-prefixed
# names (e.g. 'google/gemini-3.5-flash') and raw model ids (e.g. 'gemini-3.5-flash').
MODEL_RPM_LIMITS = {
    "google/gemini-3.5-flash": 5,
    "google/gemini-3.6-flash": 5,
    "google/gemini-3.5-flash-lite": 15,
    "google/gemini-3.1-flash-lite": 15,
    "gemini-3.5-flash": 5,
    "gemini-3.6-flash": 5,
    "gemini-3.5-flash-lite": 15,
    "gemini-3.1-flash-lite": 15,
}

# Per-model RPD limits (requests per day), confirmed against the live
# aistudio.google.com/rate-limit dashboard for this project. CRITICAL: this
# must be checked PER MODEL — previously a single hardcoded 20/day cap was
# applied to every model regardless of selection, which incorrectly blocked
# the 500 RPD Flash Lite models after only 20 total calls across ALL models.
MODEL_RPD_LIMITS = {
    "google/gemini-3.5-flash": 20,
    "google/gemini-3.6-flash": 20,
    "google/gemini-3.5-flash-lite": 500,
    "google/gemini-3.1-flash-lite": 500,
    "gemini-3.5-flash": 20,
    "gemini-3.6-flash": 20,
    "gemini-3.5-flash-lite": 500,
    "gemini-3.1-flash-lite": 500,
}

# The extraction contract the model must follow for every resume.
# Kept as a template so `{source_file}` and `{resume_text}` can be injected
# per uploaded file.
EXTRACTION_PROMPT_TEMPLATE = """
You are an advanced Resume Information Extraction Engine specialized in extracting structured data from Faculty CVs, Academic Resumes, Research Profiles, and Industry Resumes.

## Objective

Your task is to extract COMPLETE and ACCURATE information from the uploaded resume below and organize the extracted information into TWO separate logical datasets that will later be exported into TWO Excel sheets.

The output MUST contain ONLY valid JSON.

Do NOT provide explanations, markdown, comments, or extra text.

This call covers exactly ONE resume, identified as source_file = "{source_file}".

------------------------------------------------------------
OUTPUT DATASET 1 : MASTER FACULTY DATABASE
------------------------------------------------------------

This dataset represents the first Excel sheet.

This resume must generate EXACTLY ONE object inside "master_faculty_database".

Extract all available information including:

- Personal Details
- Contact Information
- Current Designation
- Department
- Organization
- Educational Qualifications
- Professional Experience
- Administrative Experience
- Research Experience
- Projects
- Patents
- Books
- Book Chapters
- Journal Counts
- Conference Counts
- Awards
- Memberships
- FDP/STTP
- Certifications
- ORCID
- Google Scholar
- ResearchGate
- Scopus ID
- LinkedIn
- Any other professional profile

If information is unavailable return null.

Never guess.

------------------------------------------------------------
OUTPUT DATASET 2 : PUBLICATION DETAILS
------------------------------------------------------------

This dataset represents the second Excel sheet.

Extract EVERY publication found in this resume.

Each publication must become ONE separate object inside "publication_details".

Never combine multiple publications.

Never omit publications.

Search the ENTIRE resume before generating the output.

Publications may appear under headings like

- Publications
- Research Publications
- Journal Papers
- Conference Papers
- Book Chapters
- Books
- Patents
- Research Output
- Scholarly Articles
- Selected Publications
- International Conference
- National Conference

They may also appear as

- Tables
- Numbered lists
- Bullet points
- Paragraphs
- Multi-line entries

Extract all publication details whenever available.

------------------------------------------------------------
CRITICAL ANTI-HALLUCINATION RULES
------------------------------------------------------------

- Do not infer, guess, or supplement information from outside the document.
- Only extract facts explicitly stated in the document text. Do not use your own general knowledge to fill in gaps (e.g., knowing which university a college is affiliated with, or expanding an abbreviation).
- If a field is not explicitly mentioned, leave it blank/null. Do not guess a plausible value.
- Do NOT infer an unstated undergraduate specialization/branch from a later degree's subject.
  Example: if the resume states "B.A." with no subject, and later states "M.A. in English" or
  "Ph.D. in English", do NOT populate education.ug.branch as "English." Leave it null. The subject
  of a later degree is never sufficient evidence for an earlier, unstated one.
- Do not compute or derive values not explicitly present, such as total years of experience, unless the document explicitly states a total (e.g., "15 years of experience"). If only start/end dates are given without an explicit total, leave the derived field blank rather than calculating it yourself.
  IMPORTANT: this rule applies ONLY to the "experience" object's academic_years/industry_years/
  research_years/administrative_years/total_years fields. It does NOT apply to designation_history —
  you must still extract "title", "organization", "date_range", and "category" for every employment
  entry stated in the resume exactly as written (e.g., "2008-Present", "2010-2015"). Extracting a
  stated date range is direct extraction, not derivation; leaving it blank when it's plainly stated
  in the document is itself an extraction error, since a separate downstream process (not you)
  computes the years-of-experience totals from these fields.
- For counts (e.g., number of projects, publications, awards): count only items unambiguously listed under the relevant section heading in the document. Do not combine, merge, or infer counts across differently-named sections.
- When a document has clearly separated table sections (e.g., "Books & Monographs" vs "Book Chapters"), classify strictly by which section the entry physically appears under. Do not rely on summary sentences elsewhere that might contradict the tables.
- Do not let adjacent column text (e.g., journal names, footnote numbers) bleed into the title or volume fields. Extract each field only from its designated column/position. If a number's role (volume vs. footnote marker) is ambiguous, leave it blank.
- If uncertain whether information is explicitly stated vs inferred, treat it as inferred and omit it.
- Your extraction will be audited. Any field containing unverified information will be considered an error.

------------------------------------------------------------
GENERAL EXTRACTION RULES
------------------------------------------------------------

1. Search the complete resume from first page to last page.
2. Never stop after finding the first publication section.
3. Preserve original spellings.
4. Preserve capitalization.
5. Never hallucinate.
6. Never fabricate missing values.
7. Return null when information is unavailable.
8. Ignore duplicate publications.
9. Calculate experience only if employment dates are available.
10. Preserve author order.
11. Preserve publication titles exactly.
12. If publication numbering exists, use it.
13. Otherwise generate sequential numbering.
14. Extract all emails belonging to the CV OWNER only. Many CVs include a "References," "Referees," or similar section near the end listing OTHER people (e.g., supervisors, colleagues, department heads) along with their own emails and phone numbers as character references. These belong to the referees, NOT the CV owner — never include a reference/referee's email in the owner's "emails" field, even though it appears in the same document. Only extract emails that are presented as the CV owner's own contact details (typically in the header, contact section, or personal profile section).
15. Extract all phone numbers belonging to the CV OWNER only, with the same exclusion of reference/referee phone numbers as rule 14. Additionally, if the same phone number appears more than once in the document written in different formats (e.g., "+91-9014516726" and "9014516726" — the same digits with/without a country code or punctuation), include it only ONCE in the phones array, using the more complete formatted version (the one with the country code, if given). Do not list the same underlying number twice just because it was typed differently in two places.
16. Choose the latest organization as current organization.
17. Choose the latest designation as current designation.
18. If publication type is unclear classify using:
    Journal, International Conference, National Conference, Book, Book Chapter, Patent, Other
19. Extract DOI, ISSN, ISBN, Quartile, Impact Factor, Indexing, URL whenever available.
20. For designation_history, set category to "academic" for teaching ranks, "industry" for company roles, or "research" for Research Associate/Scientist/Fellow.
    "category" and "date_range" are REQUIRED for every designation_history entry whenever the
    resume states them — do not omit these two fields even though most other fields are optional.
    IMPORTANT: Never emit a `date_range` that ends with a bare dash or connector (e.g. "2010-", "May 2010 –").
    For ongoing/current roles explicitly write the end as "Present" (examples: "2010-Present", "May 2010 – Present").
    Do NOT use a trailing dash to indicate continuity; the downstream parser treats a bare-dash as incomplete.
    Downstream code computes years of experience purely from these two fields, so leaving them
    blank when the information is present in the resume is itself an extraction error.
21. If the resume has a single undifferentiated 'Projects' section rather than separate Sponsored/Consultancy/Training headings, put entries in projects_list instead of guessing a category.
22. Populate the 'flags' array when you deliberately omit an ambiguous field (e.g., ["ambiguous_total_experience", "unclear_pub_type"]).
23. STRICT DISAMBIGUATION — projects vs. trainings vs. attended FDPs/STTPs:
    - "sponsored_projects" / "consultancy_projects" / "projects_list" = funded research work,
      consultancy assignments, or industry-sponsored projects with a sponsor/client named.
    - "trainings_conducted" = FDPs, workshops, STTPs, or orientation/refresher programmes that
      THIS PERSON ORGANIZED, DELIVERED, OR TAUGHT (i.e. they were the resource person/organizer/
      coordinator, not a participant). Look for verbs like "conducted", "organized", "delivered a
      session", "resource person for".
    - "fdp_list" = Faculty Development Programmes this person ATTENDED/PARTICIPATED IN as a
      learner (look for verbs like "attended", "participated in", "completed"). Do NOT put these
      in trainings_conducted.
    - "sttp_list" = Short-Term Training Programmes (STTPs) this person ATTENDED/PARTICIPATED IN
      as a learner, same distinction as fdp_list above.
    - These are never projects, even if the resume calls them a "programme" or lists them near a
      projects section.
    - The distinguishing question for every FDP/workshop/STTP entry is: did this person run it, or
      attend it? Conducted/organized/delivered → trainings_conducted. Attended/participated →
      fdp_list or sttp_list depending on which the entry names. If the resume lists a combined
      "Workshops/FDPs Attended" section, every entry in it goes to fdp_list, not trainings_conducted,
      regardless of whether a nearby heading elsewhere also contains the word "conducted".
24. Return ONLY valid JSON, matching the schema below exactly.
25. Every object in "publication_details" must include "faculty_name" and "source_file" matching this resume.
26. Do not create separate education.*.university or education.*.institute fields. Use the single combined education.*.institution field for all education entries.
27. For education fields, if a resume mentions both a university name and an institute/college name for the same degree, combine them into the single 'institution' field separated by a comma (e.g., 'ABC College of Engineering, XYZ University').
28. "education.pg" is an ARRAY, not a single object. If a resume lists more than one postgraduate-level degree (e.g., an MCA earned before an M.Tech, or an M.Phil earned between a Master's and a PhD), include EACH one as its own object in the pg array, in chronological order. Never drop an intermediate postgraduate degree just because a later, more advanced one also exists.
29. Common postgraduate-level degree abbreviations (M.Sc, M.A., M.Com, M.Tech, MBA, M.Phil, ME, MS, MCA) must NEVER be placed in education.ug — these always belong in education.pg. If you are unsure whether a degree is UG or PG level, use its full name/context in the resume to decide; do not guess.
30. If the resume contains a standalone SUMMARY/AGGREGATE TABLE for publications (e.g., a table or line stating "International Journal: 99, International Conference: 70, National Conference: 11, Books: 7, Book Chapters: 5") that is separate from — or in addition to — an itemized list of individual publications, extract these aggregate figures directly into the "counts" object (journal, international_conference, national_conference, books, book_chapters) even if you cannot also produce an individual publication_details object for every one of those counted items. Do not skip this table just because you already found an itemized list elsewhere — some resumes have both, and the summary table's figures should be preferred if the itemized list appears incomplete relative to it. If you use a summary table's figures instead of (or in addition to) an itemized list, add the flag "counts_from_summary_table".
31. Competitive/qualifying exam results (e.g., "Qualified UGC-NET", "Qualified GATE", "Qualified SET", "Cleared JRF") that are listed under an Achievements, Awards, or Honors heading count as entries in awards_list — do not exclude them just because they are exam qualifications rather than named prizes.
32. Before finalizing awards_list, check for near-duplicate entries describing the SAME event/medal/award with only minor wording differences (e.g., "Gold Medal Inter IIT Sport Meet 2018" and "Gold Medal in IIT Sports Meet 2018 in Para-powerlifting" likely describe one award). If two entries clearly refer to the same underlying award, include it only ONCE, using the more descriptive/complete wording.
32b. GENERAL EXHAUSTIVENESS RULE — applies to every list field in this schema (awards_list, memberships_list, administrative_roles, workshops_attended_list, online_courses_list, sponsored_projects, consultancy_projects, trainings_conducted, publication_details, and any other array): when a resume section contains a long bulleted or numbered list (10+ items), you MUST process and include every single line item individually. Do not stop early, do not summarize or paraphrase the list into a shorter version, and do not skip items merely because they look similar to ones already included (only merge near-duplicates per the specific dedup rules already given for awards/publications). Long lists are exactly where omissions happen most under time/token pressure — treat completeness on a 30-item list as being just as important as completeness on a 3-item list. If you are ever tempted to write "and other similar items" or stop partway through a long list, that impulse is the signal to keep going instead.
32c. For administrative_roles specifically: if a section heading or nearby text establishes a date or date range that applies to the whole list (e.g., "Administrative Work (2018-2023):" or a role described elsewhere in the same resume as "since June 2021"), attach that date_range to every role in that list that shares the same time period, rather than leaving date_range blank for items that do have an inferable date from context. Only leave date_range blank when the resume gives genuinely no date information anywhere near that item.
33. If the resume has separate sections for "Workshops Attended", "Seminars/Conferences Participated (as attendee, not presenter)", or "Online Courses/MOOCs Completed" (e.g., NPTEL, Coursera, edX), extract these into two new arrays: "workshops_attended_list" (workshops and seminars attended as a participant) and "online_courses_list" (any completed online course/MOOC). These are distinct from fdp_list/sttp_list (formal Faculty Development Programmes/Short-Term Training Programmes) and from trainings_conducted (things the person organized/delivered) — use them only for generic workshops/seminars/online courses that don't fit those more specific buckets.
34. If the resume explicitly states a TOTAL years-of-experience figure anywhere (e.g., in a header, summary, or objective section — such as "12 years 5 months of experience" or "Total 22 years of teaching experience"), extract it VERBATIM as a string into "experience.stated_total_years". This is in ADDITION to (not instead of) extracting designation_history entries normally — downstream code computes its own total from date ranges, and stated_total_years lets that computed figure be compared against what the resume itself claims.
35. If the resume states an aggregate LIFETIME publication count (e.g., "240 publications", "over 100 research papers") that is significantly higher than the number of individual publication_details items you were able to extract from the visible list/table, add the flag "possible_incomplete_publication_list_vs_stated_total" to the master record. Do NOT fabricate additional publication_details entries to make the count match — only flag the discrepancy.

------------------------------------------------------------
JSON OUTPUT FORMAT
------------------------------------------------------------

{{
  "master_faculty_database": [
    {{
      "source_file": "{source_file}",
      "name": "",
      "emails": [],
      "phones": [],
      "nationality": "",
      "country": "",
      "address": "",
      "current_designation": "",
      "current_department": "",
      "current_organization": "",
      "education": {{
        "ug": {{"degree": "", "branch": "", "institution": "", "year": ""}},
        "pg": [
          {{"degree": "", "branch": "", "institution": "", "year": ""}}
        ],
        "phd": {{"degree": "", "branch": "", "institution": "", "year": ""}}
      }},
      "designation_history": [
        {{"title": "", "organization": "", "date_range": "", "category": ""}}
      ],
      "administrative_roles": [
        {{"title": "", "date_range": ""}}
      ],
      "sponsored_projects": [],
      "consultancy_projects": [],
      "trainings_conducted": [],
      "projects_list": [],
      "awards_list": [],
      "memberships_list": [],
      "fdp_list": [],
      "sttp_list": [],
      "workshops_attended_list": [],
      "online_courses_list": [],
      "experience": {{
        "academic_years": null,
        "industry_years": null,
        "research_years": null,
        "administrative_years": null,
        "total_years": null,
        "stated_total_years": ""
      }},
      "counts": {{
        "journal": 0,
        "international_conference": 0,
        "national_conference": 0,
        "books": 0,
        "book_chapters": 0,
        "patents": 0,
        "projects": 0,
        "sponsored_projects": 0,
        "consultancy_projects": 0,
        "trainings": 0,
        "projects_general": 0,
        "awards": 0,
        "memberships": 0,
        "fdp_attended": 0,
        "sttp_attended": 0,
        "workshops_attended": 0,
        "online_courses": 0
      }},
      "flags": [],
      "profiles": {{
        "orcid": "",
        "google_scholar": "",
        "researchgate": "",
        "linkedin": "",
        "scopus": ""
      }}
    }}
  ],
  "publication_details": [
    {{
      "faculty_name": "",
      "source_file": "{source_file}",
      "sr_no": 1,
      "title": "",
      "authors": [],
      "publication_type": "",
      "journal_or_conference": "",
      "publisher": "",
      "volume": "",
      "issue": "",
      "pages": "",
      "year": "",
      "month": "",
      "doi": "",
      "issn": "",
      "isbn": "",
      "indexed_in": [],
      "quartile": "",
      "impact_factor": "",
      "citation_count": null,
      "url": ""
    }}
  ]
}}

------------------------------------------------------------
RESUME TEXT (source_file = "{source_file}")
------------------------------------------------------------

{resume_text}

------------------------------------------------------------
FINAL INSTRUCTION
------------------------------------------------------------

Return ONLY the JSON object described above. No markdown, no code fences, no commentary.
"""


class ResumeAnalyzerEngine:
    def __init__(self):
        # Configure Google Studio API key and model
        self.api_key = os.getenv("GOOGLE_STUDIO_API_KEY")
        if not self.api_key:
            logger.warning("GOOGLE_STUDIO_API_KEY is not set in environment variables.")
        self.base_url = config.GOOGLE_STUDIO_BASE_URL
        self.model_name = config.GOOGLE_STUDIO_MODEL

        logger.info(f"Initialized ResumeAnalyzerEngine with model: {self.model_name}")

        # Confirmed against the live AI Studio dashboard for this project
        # (aistudio.google.com/rate-limit) rather than assumed from generic
        # blog figures: gemini-3.5-flash free tier here is 5 RPM / 20 RPD,
        # not the 20 RPM this code originally (incorrectly) assumed. RPD is
        # the tighter constraint by far — verify your own account's current
        # numbers there if you upgrade tiers or switch models, since these
        # are not fixed platform-wide constants.
        # Per-model timestamp queues to enforce model-specific RPM.
        # Key: normalized model string; Value: deque of recent timestamps (seconds)
        self._request_timestamps_by_model = {}
        self._rate_limit_lock = threading.Lock()
        # Fallback RPM/RPD used when a model is not listed in the limits dicts
        self._default_rpm = 5
        self._default_rpd = 20
        # Per-model daily request tracking: key -> deque of dates (Pacific)
        # of recent calls FOR THAT MODEL SPECIFICALLY. Previously this was a
        # single shared deque with a hardcoded 20/day cap applied to every
        # model, which incorrectly throttled 500 RPD models after only 20
        # total calls across all models combined.
        self._daily_request_dates_by_model = {}

    def _get_model_rpd(self, model_name: str) -> int:
        key = self._get_normalized_model_key(model_name)
        return MODEL_RPD_LIMITS.get(key, self._default_rpd)

    def google_studio_daily_usage(self, model_name: str = None):
        """Returns (used_today, limit) for the given model's free-tier RPD,
        based on this process's own count of requests it has sent today FOR
        THAT MODEL. This is a best-effort local tracker, not a value fetched
        from Google — the API itself doesn't expose remaining quota. It
        resets whenever the app restarts, and won't know about calls made
        outside this running process.
        For the authoritative live number, check aistudio.google.com/rate-limit."""
        key = self._get_normalized_model_key(model_name)
        limit = self._get_model_rpd(model_name)
        with self._rate_limit_lock:
            today = self._pacific_date_today()
            dq = self._daily_request_dates_by_model.setdefault(key, deque())
            while dq and dq[0] != today:
                dq.popleft()
            return len(dq), limit

    def _pacific_date_today(self):
        # RPD resets at midnight Pacific time, not local/UTC time.
        try:
            from zoneinfo import ZoneInfo
            import datetime
            return datetime.datetime.now(ZoneInfo("America/Los_Angeles")).date()
        except Exception:
            # Fallback if zoneinfo/tzdata isn't available: approximate with UTC-8,
            # which is wrong during PDT but avoids a hard crash.
            import datetime
            return datetime.datetime.utcfromtimestamp(time.time() - 8 * 3600).date()

    def _check_daily_budget(self, source_file: str, model_name: str = None):
        """Raises a clear, actionable error immediately if today's RPD budget
        for THIS SPECIFIC MODEL is already spent, instead of retrying against
        a quota that won't reset for hours. Each model has its own daily
        counter, since the models have very different real RPD limits
        (20/day vs 500/day) that must not share a single global counter."""
        key = self._get_normalized_model_key(model_name)
        limit = self._get_model_rpd(model_name)
        with self._rate_limit_lock:
            today = self._pacific_date_today()
            dq = self._daily_request_dates_by_model.setdefault(key, deque())
            while dq and dq[0] != today:
                dq.popleft()
            used = len(dq)
            if used >= limit:
                return {
                    "error": (
                        f"'{source_file}': Google Studio free-tier daily quota for this model "
                        f"({limit} requests/day) is exhausted for today "
                        f"({used}/{limit} used). This resets at midnight "
                        f"Pacific time (~12:30 PM IST). To process more today, switch to a "
                        f"higher-quota model, enable billing on this Gemini project "
                        f"(aistudio.google.com/apikey) for Tier 1 limits, or retry after the daily reset."
                    )
                }
            dq.append(today)
            return None

    def _get_normalized_model_key(self, model_name: str) -> str:
        if not model_name:
            return ""
        m = model_name.lower()
        return m

    def _get_model_rpm(self, model_name: str) -> int:
        key = self._get_normalized_model_key(model_name)
        return MODEL_RPM_LIMITS.get(key, self._default_rpm)

    def get_model_wait_seconds(self, model_name: str) -> float:
        """Return a conservative per-request interval (seconds) for the given model.
        Caller-facing helper used by the UI/driver to pace between-file processing.
        """
        rpm = self._get_model_rpm(model_name)
        # Add a small buffer to avoid edge collisions
        return 60.0 / rpm + 0.5

    def _wait_for_rate_limit_slot(self, model_name: str = None):
        """Block until a rate-limit slot is available for the specified model.
        This enforces per-model RPM guarantees using a sliding 60-second window.
        """
        key = self._get_normalized_model_key(model_name)
        rpm = self._get_model_rpm(model_name)
        with self._rate_limit_lock:
            if key not in self._request_timestamps_by_model:
                self._request_timestamps_by_model[key] = deque()
            dq = self._request_timestamps_by_model[key]

        while True:
            with self._rate_limit_lock:
                now = time.time()
                # Purge entries older than 60s
                while dq and now - dq[0] > 60:
                    dq.popleft()
                if len(dq) < rpm:
                    dq.append(now)
                    return
                sleep_for = 60 - (now - dq[0]) + 0.5
            logger.info(f"Rate limit reached for model {model_name} ({rpm}/min); waiting {sleep_for:.1f}s for a slot.")
            time.sleep(max(sleep_for, 0.5))

    def extract_text_from_pdf(self, pdf_path: str, progress_callback=None) -> str:
        """Extract text from a PDF file using pdfplumber, falling back to OCR if empty.

        progress_callback: optional callable(str) invoked with a short status
        string after each page/step, so the caller (e.g. Streamlit) can show
        live progress instead of a frozen spinner.
        """
        def _report(msg):
            logger.info(msg)
            if progress_callback:
                try:
                    progress_callback(msg)
                except Exception:
                    pass  # never let a UI callback break extraction

        text = ""
        try:
            with pdfplumber.open(pdf_path) as pdf:
                total_pages = len(pdf.pages)
                _report(f"Opened PDF with {total_pages} page(s). Trying direct text extraction...")
                for i, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
                    _report(f"Direct extraction: page {i}/{total_pages} done "
                             f"({len(page_text) if page_text else 0} chars).")

            if text.strip():
                _report("Direct PDF text extraction succeeded.")
                return text.strip()
            else:
                _report("Direct extraction produced no text — this looks like a scanned/image PDF.")
        except Exception as e:
            logger.error(f"Direct text extraction failed: {e}")
            _report(f"Direct text extraction raised an error: {e}")

        # Fallback to OCR for scanned/image-based PDFs
        _report("Falling back to OCR (pytesseract). Converting pages to images...")
        try:
            images = convert_from_path(pdf_path, dpi=config.OCR_DPI)
            total_images = len(images)
            _report(f"Converted to {total_images} image(s) at {config.OCR_DPI} DPI. Starting OCR...")
            
            if total_images > config.OCR_MAX_PAGES:
                _report(f"WARNING: PDF has {total_images} pages. Capping OCR at {config.OCR_MAX_PAGES} pages to prevent timeouts.")
                images = images[:config.OCR_MAX_PAGES]
                
            for i, image in enumerate(images, start=1):
                _report(f"OCR: processing page {i}/{len(images)} (of {total_images} total)...")
                page_text = pytesseract.image_to_string(image)
                text += page_text + "\n"
                _report(f"OCR: page {i}/{len(images)} done "
                         f"({len(page_text)} chars extracted).")
        except Exception as e:
            logger.error(f"OCR text extraction failed: {e}")
            _report(f"OCR failed: {e}")

        return text.strip()

    def extract_text_from_docx(self, docx_path: str, progress_callback=None) -> str:
        """Extract text from a DOCX file by reading paragraphs and table cells."""
        def _report(msg):
            logger.info(msg)
            if progress_callback:
                try:
                    progress_callback(msg)
                except Exception:
                    pass

        text_parts = []
        try:
            document = docx.Document(docx_path)
            _report("Opened DOCX document. Extracting paragraphs and tables...")

            for i, para in enumerate(document.paragraphs, start=1):
                if para.text and para.text.strip():
                    text_parts.append(para.text.strip())
                if i % 20 == 0:
                    _report(f"DOCX extraction: processed {i} paragraphs...")

            for table_index, table in enumerate(document.tables, start=1):
                _report(f"DOCX extraction: processing table {table_index}/{len(document.tables)}...")
                for row in table.rows:
                    row_text = " \t ".join(cell.text.strip() for cell in row.cells if cell.text and cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            extracted_text = "\n".join(text_parts).strip()
            if extracted_text:
                _report("DOCX text extraction succeeded.")
            else:
                _report("DOCX text extraction produced no readable text.")
            return extracted_text
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            _report(f"DOCX extraction failed: {e}")
            return ""

    @staticmethod
    def _clean_json_text(raw_text: str) -> str:
        """Strip markdown code fences if the model adds them despite instructions."""
        cleaned = raw_text.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.split("```", 2)[1] if cleaned.count("```") >= 2 else cleaned.strip("`")
            cleaned = cleaned[4:] if cleaned.lower().startswith("json") else cleaned
            cleaned = cleaned.strip()

        # Some models (especially smaller open-weight ones) prepend commentary
        # before the JSON object. Extract from the first '{' to the matching
        # last '}' rather than trusting the whole string is valid JSON.
        first_brace = cleaned.find("{")
        last_brace = cleaned.rfind("}")
        if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
            cleaned = cleaned[first_brace:last_brace + 1]

        return cleaned.strip()

    def _prepare_google_studio_prompt(self, source_file: str, resume_text: str) -> str:
        """Trim the search prompt to stay under Google Studio's request payload limit."""
        prompt = EXTRACTION_PROMPT_TEMPLATE.format(
            source_file=source_file,
            resume_text=resume_text,
        )

        def request_size_for_text(text: str) -> int:
            # Mirrors the actual payload shape sent in extract_faculty_data
            # (Gemini's generateContent contents/parts format) so this size
            # estimate is accurate.
            payload = {
                "contents": [
                    {"parts": [{"text": text}]}
                ],
                "generationConfig": {
                    "temperature": config.EXTRACTION_TEMPERATURE,
                    "topP": 1,
                    "topK": 1,
                    "seed": 42,
                    "maxOutputTokens": 32768,
                    "responseMimeType": "application/json",
                },
            }
            return len(json.dumps(payload, ensure_ascii=False).encode("utf-8"))

        max_bytes = config.GOOGLE_STUDIO_MAX_REQUEST_BYTES
        if request_size_for_text(prompt) <= max_bytes:
            return prompt

        # Binary search for the largest resume_text slice that fits in the request.
        low, high = 0, len(resume_text)
        fitted_text = ""
        while low <= high:
            mid = (low + high) // 2
            candidate_text = resume_text[:mid]
            candidate_prompt = EXTRACTION_PROMPT_TEMPLATE.format(
                source_file=source_file,
                resume_text=candidate_text,
            )
            if request_size_for_text(candidate_prompt) <= max_bytes:
                fitted_text = candidate_text
                low = mid + 1
            else:
                high = mid - 1

        if not fitted_text:
            logger.warning(
                "Google Studio prompt was too large even with empty resume text, falling back to a minimal request."
            )
            return EXTRACTION_PROMPT_TEMPLATE.format(source_file=source_file, resume_text="")

        logger.warning(
            "Google Studio request payload exceeded %s bytes and was truncated to fit. "
            "This may reduce extraction coverage.",
            max_bytes,
        )
        return EXTRACTION_PROMPT_TEMPLATE.format(
            source_file=source_file,
            resume_text=fitted_text,
        )

    def extract_faculty_data(self, resume_text: str, source_file: str, model_name: str = None) -> dict:
        """
        Call the configured Google Studio model to extract structured Master Faculty Database and
        Publication Details records for a single resume.

        Returns a dict: {"master_faculty_database": [...], "publication_details": [...]}
        or {"error": "..."} on failure.
        """
        if not resume_text or not resume_text.strip():
            return {"error": f"'{source_file}': resume text is empty or could not be extracted."}

        actual_model = model_name if model_name else self.model_name
        provider = "Google Studio"
        req_base_url = self.base_url
        req_api_key = self.api_key

        prompt = self._prepare_google_studio_prompt(source_file, resume_text)

        if actual_model.lower().startswith("google/"):
            actual_model = actual_model.split("/", 1)[1]

        headers = {"Content-Type": "application/json"}
        url = f"{req_base_url}/models/{actual_model}:generateContent"
        if req_api_key:
            url = f"{url}?key={req_api_key}"
        payload = {
            "contents": [
                {"parts": [{"text": prompt}]}
            ],
            "generationConfig": {
                "temperature": config.EXTRACTION_TEMPERATURE,
                "topP": 1,
                "topK": 1,
                "seed": 42,
                "maxOutputTokens": 32768,
                "responseMimeType": "application/json",
            },
        }

        try:
            budget_error = self._check_daily_budget(source_file, actual_model)
            if budget_error:
                return budget_error
            # Log which model string is actually used for this call (helps debug dropdown selection issues)
            logger.info(f"Using model for API call: {actual_model} (source_file={source_file})")
            self._wait_for_rate_limit_slot(actual_model)

            # Only 2 attempts for Google Studio: a 429 might be the 5 RPM cap
            # (worth a short retry) or the 20 RPD cap (won't resolve for
            # hours) — we can't always tell which from the response, so we
            # retry once rather than burning several attempts against a
            # quota that's actually exhausted for the day.
            max_attempts = 2
            response = None
            for attempt in range(1, max_attempts + 1):
                response = requests.post(url, headers=headers, json=payload, timeout=180)
                if response.status_code != 429:
                    break
                if attempt == max_attempts:
                    break
                # Google's 429 body includes a human-readable "Please retry in
                # X.Ys" hint — honor it exactly instead of guessing a backoff.
                retry_after = None
                match = re.search(r"retry in ([\d.]+)s", response.text)
                if match:
                    retry_after = float(match.group(1))
                wait_s = (retry_after if retry_after is not None else 20.0) + 1.0
                logger.warning(
                    f"429 from {provider} for {source_file} (attempt {attempt}/{max_attempts}); "
                    f"retrying in {wait_s:.1f}s."
                )
                time.sleep(wait_s)
                self._wait_for_rate_limit_slot(actual_model)

            try:
                body = response.json()
            except json.JSONDecodeError:
                response.raise_for_status()
                raise

            if response.status_code >= 400:
                err_msg = None
                if isinstance(body, dict) and "error" in body:
                    err = body["error"]
                    if isinstance(err, dict):
                        err_msg = err.get("message") or err.get("code") or str(err)
                    else:
                        err_msg = str(err)
                if not err_msg:
                    err_msg = body.get("message") if isinstance(body, dict) else response.text
                logger.error(f"{provider} returned HTTP {response.status_code} for {source_file}: {err_msg}")
                return {"error": f"'{source_file}': {provider} error: {err_msg}"}

            if "error" in body:
                err_msg = body["error"].get("message", str(body["error"])) if isinstance(body["error"], dict) else str(body["error"])
                logger.error(f"{provider} returned an error payload for {source_file}: {err_msg}")
                return {"error": f"'{source_file}': {provider} error: {err_msg}"}

            candidates = body.get("candidates") or []
            if not candidates:
                logger.error(f"Unexpected Google Studio response for {source_file}: {body}")
                return {"error": f"'{source_file}': model returned no output. Try again shortly or verify the model name and API key."}
            first_candidate = candidates[0]
            # generateContent's real response shape:
            # {"candidates": [{"content": {"parts": [{"text": "..."}]}}]}
            content = first_candidate.get("content") or {}
            parts = content.get("parts") or []
            raw_text = "".join(p.get("text", "") for p in parts if isinstance(p, dict))
            finish_reason = first_candidate.get("finishReason")

            # Log exact per-call cost from Gemini's real usageMetadata
            # instead of estimating. Thinking tokens are billed as output
            # tokens, so thoughtsTokenCount is included in the cost calc.
            usage = body.get("usageMetadata") or {}
            if usage:
                prompt_tokens = usage.get("promptTokenCount", 0)
                output_tokens = usage.get("candidatesTokenCount", 0)
                thought_tokens = usage.get("thoughtsTokenCount", 0)
                total_output_tokens = output_tokens + thought_tokens
                # gemini-3.5-flash list price as of 2026-07: $1.50/M input,
                # $9.00/M output (thinking tokens billed at output rate).
                # Update these if you switch models or Google repriced.
                cost = (prompt_tokens * 1.50 + total_output_tokens * 9.00) / 1_000_000
                logger.info(
                    f"Token usage for {source_file}: prompt={prompt_tokens}, "
                    f"output={output_tokens}, thinking={thought_tokens}, "
                    f"total_output={total_output_tokens} -> est. cost ${cost:.5f}"
                )

            if finish_reason == "MAX_TOKENS":
                # The response was cut off mid-generation — raw_text may
                # be non-empty but is guaranteed-invalid partial JSON.
                # Fail clearly here rather than passing it to json.loads()
                # and surfacing a confusing "Expecting ',' delimiter" error.
                logger.error(
                    f"Google Studio response for {source_file} hit MAX_TOKENS "
                    f"({len(raw_text)} chars generated before cutoff)."
                )
                return {
                    "error": f"'{source_file}': model response was cut off (hit the output "
                    "token limit) before finishing the JSON. This resume likely has a lot of "
                    "publications/content — try again, or reduce resume length."
                }

            if not raw_text:
                logger.error(
                    f"Empty text from Google Studio for {source_file}. "
                    f"finishReason={finish_reason}. Full candidate: {first_candidate}"
                )
                return {"error": f"'{source_file}': model returned empty output. Try again, or switch to a different model in the sidebar."}

            cleaned = self._clean_json_text(raw_text)
            data = json.loads(cleaned)

            master_records = data.get("master_faculty_database", []) or []
            publication_records = data.get("publication_details", []) or []
            
            publication_records = dedupe_publications(publication_records)
            publication_records = renumber_publications(publication_records)

            # Enforce source_file consistency and backfill faculty_name on
            # publications so downstream aggregation/export never has to guess.
            faculty_name = None
            for i, record in enumerate(master_records):
                # Always enforce correct source_file in records
                record["source_file"] = source_file
                faculty_name = record.get("name") or faculty_name
                master_records[i] = safe_process_record(record, publication_records)

            for pub in publication_records:
                pub["source_file"] = source_file
                if not pub.get("faculty_name"):
                    pub["faculty_name"] = faculty_name

            return {
                "master_faculty_database": master_records,
                "publication_details": publication_records,
            }

        except requests.exceptions.HTTPError as he:
            logger.error(f"{provider} API HTTP error for {source_file}: {he}")
            return {"error": f"'{source_file}': {provider} API error: {str(he)}"}
        except json.JSONDecodeError as jde:
            logger.error(f"Failed to parse JSON from {provider} response for {source_file}: {jde}")
            # raw_text might not exist if it failed before assignment, but in this scope it does
            try:
                logger.error(f"Raw model output (first 1000 chars): {raw_text[:1000]}")
            except Exception:
                pass
            return {"error": f"'{source_file}': received malformed extraction output from the model. Please try again."}
